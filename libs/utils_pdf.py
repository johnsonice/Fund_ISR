### extract ram table from AIVs 
#%%
import PyPDF2
import re
from docx import Document
#%%
# Save the resulting PDF to a file.
def get_page(filename,page_num):
    pdfFileObj = open(filename,'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    pageObj = pdfReader.getPage(page_num)
    content = pageObj.extract_text().lower().replace('\n','').replace('  ',' ')
    pdfFileObj.close()
    return content

def key_reg_match(content,reg_text=None,search_text=None,mode='all'):
    """use either regular expressions or keywords matching"""
    if reg_text:
        if mode in ['all','rgx']:
            result = reg_text.findall(content)
            if len(result)>0:
                return True
            
    if search_text:
        if mode in ['all','key']:
            if search_text in content:
                return True
    ## if nothing matched     
    return None

#filename = 'U:/STA/documents/Bangladesh 2017.pdf'
#search_text = 'table of common indicators'# required for surveillance'
def get_target_pagenum(filename,match_func,return_mode='first'): ## mode = first or all
    pdfFileObj = open(filename,'rb')
    pdfReader = PyPDF2.PdfReader(pdfFileObj)
    all_res = [] 
    for pageNum in range(len(pdfReader.pages)):
            pageObj = pdfReader.pages[pageNum]
            content = pageObj.extract_text().lower().replace('\n','').replace('  ',' ')
            match_result = match_func(content)
            if match_result:
                if return_mode=='first':
                    pdfFileObj.close()
                    return pageNum
                if return_mode=='all':
                    all_res.append(pageNum)
            
    pdfFileObj.close()
    return all_res
class Word_Table_Extractor():
    def __init__(self, filepath=None):
        self.filepath = filepath
        self.is_valid_table = None
        if filepath:
            self.load_doc(filepath)
        
    def load_doc(self,filepath):
        self.doc = Document(filepath)
        self.tables = self.doc.tables 
        if len(self.tables):
            self.t0 = self.tables[0]
        else:
            self.t0 = None
            
    def clear_doc(self):
        self.doc = None
        self.tables = None
        self.t0 = None
        
    @staticmethod
    def basic_table_extract(table,headers=None):
        data = []
        # Extract column headers
        if headers:
            pass
        else:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Extract row data
        for row in table.rows[1:]:
            row_data = {headers[i]: row.cells[i].text.strip() for i in range(len(row.cells))}  ## there is a problem with using dict type: duplicated headers will be replaced ; fix later 
            data.append(row_data)
            
        return data

class Word_Table_Identifier(Word_Table_Extractor):
    def __init__(self, filepath=None):
        super().__init__(filepath)

    @staticmethod
    def _remove_text_within_parentheses(text):
        pattern = r"\(.*?\)"
        return re.sub(pattern, '', text).strip().replace("\n"," ") ## also replace ]n
    
    @staticmethod
    def _check_row_column(t0,check_type='row',check_n=3):
        """
        check if there are merged rows or cells 
        """
        ## make sure do not over subscribe 
        if check_type == 'row':
            len_check_item = len(t0.rows)
            check_cells = t0.row_cells
        elif check_type == 'column':
            len_check_item = len(t0.columns)
            check_cells = t0.column_cells
        if not len_check_item>check_n: check_n=len_check_item ## check check to <= max len
        # first filter and get cells that has longer than 7 words 
        ## if string is relatively long, they are not headers ; then they shouldn't have duplicates 
        check_items = [[c.text for c in check_cells(i) if len(c.text.strip().split())>7] for i in range(check_n)]
        for c_i in check_items:
            if len(c_i) != len(set(c_i)):
                return True ## identified dups 
        return False ## no duplicates or merged cells 

    ## rule 1 : check if column headers is valid 
    def _rule_1(self,t0,verbose=True):
        """
        - check first 1-3 row, see if any of the 3 rows roughly contains risks; likelihood
        - raw info shouldn't be too long (after removing info in "()"); sometime they can have something like this "Source of Risks and Relative Likelihood (High, medium, or low)"
        """
        if len(t0.rows)<=2:
            if verbose:
                print('rule 1 failed: table 0 less than 3 rows')
            return False
        
        for row in t0.rows[:3]:
            maybe_headers_strings = " ".join([cell.text.strip() for cell in row.cells])
            maybe_headers_strings = self._remove_text_within_parentheses(maybe_headers_strings)
            pattern = re.compile(r'(?=.*(?:risk|threat|shock|source))(?=.*(?:likelihoo|impact))', re.I) 
            res = key_reg_match(maybe_headers_strings,reg_text = pattern,mode='rgx')
            if res:
                if len(maybe_headers_strings.split()) < 30:
                    return True ## columns header identified ; pass
        if verbose:
            print('Rule 1 failed: no header identified')
        return False ## no header fund 
    
    ## rule 2 : check if if there are merged columns ; by checking if there are duplicated cells in a row ;  
    def _rule_2(self,t0,verbose=True):
        if self._check_row_column(t0,check_type='row',check_n=7): ## if identified dups
            if verbose:
                print('rule 2 failed: identified merged cells row wise... ')
            return False 
        
        if self._check_row_column(t0,check_type='column',check_n=2): ## if identified dups
            if verbose:
                print('rule 2 failed: identified merged cells column wise ... ')
            return False 
        
        return True # if no dup found, return True 
    
    ## rule 3 : sometimes, the entire row or table merged into once cell check for extremely long cell content or check for multiple \n\n\n
    def _rule_3(self,t0,verbose=True):
        ## check all cells ; shouldn't have any cell value that is strange 
        for row in t0.rows:
            for c in row.cells:
                text = c.text
                triple_newline_occurrences = text.count("\n\n\n") > 2
                # Check if the number of words is over 500
                word_count_over= len(text.split()) > 200
                if  triple_newline_occurrences or word_count_over:
                    if verbose:
                        print('rule 3 failed: large merged cell identified')
                    return False ## identified problem  cell value
        
        return True
    
    def identify_non_standard_tables(self,table=None,doc_path=None,verbose=True):
        if doc_path:
            self.load_doc(doc_path)
            if len(self.tables) == 0:
                if verbose:
                    print('rule 0 failed: no table found')
                return False
            
        if table:
            self.clear_doc()
            self.t0 = table
            
        rules = [self._rule_1,self._rule_2,self._rule_3]
        for r in rules:
            if not r(self.t0,verbose):
                return False
        
        return True
    
# def get_table_end(filename,start_page_num,check_page_n=1):
#     pdfFileObj = open(filename,'rb')
#     pdfReader = PyPDF2.PdfReader(pdfFileObj)
#     all_res = []
#     for pageNum in range(start_page_num+1,start_page_num+check_page_n+1):
#             pageObj = pdfReader.pages[pageNum]
#             content = pageObj.extract_text().lower().replace('\n','').replace('  ',' ')
#             #pattern = re.compile(r'likelihood|policy response',re.I)
#             pattern = re.compile(r'(?=.*likelihood)(?=.*policy response)(?=.*risk)')
#             res = key_reg_match(content,reg_text = pattern,mode='rgx')
#             if res:
#                 all_res.append(pageNum)
#             else:
#                 if len(all_res)>0:
#                     return all_res[-1]

# ## example of a matching function 
# def matching_function(content):
#     match_key = 'risk assessment matrix'
#     if key_reg_match(content,search_text=match_key,mode='key'):
#             pattern = re.compile(r'likelihood|policy response',re.I)
#             # Search the text for the pattern
#             res = key_reg_match(content,reg_text = pattern,mode='rgx')
#             return res 
#     return None

# #%%
# if __name__ == "__main__":
#     file_path = "/root/workspace/data/DOCs/PDF/USA_2022.pdf"
#     pgn = get_target_pagenum(file_path,matching_function,return_mode='first')
#     end_pgn = get_table_end(file_path,pgn,check_page_n=2)
#     print(pgn,end_pgn)

# #%%




